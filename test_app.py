"""
test_app.py
app.py のユニットテスト・インテグレーションテスト

実行:
    pip install pytest pytest-asyncio httpx
    pytest test_app.py -v
"""
from __future__ import annotations

import json
import shutil
import sys
import tempfile
import types
import unittest.mock as mock
from pathlib import Path
from unittest.mock import MagicMock, patch, PropertyMock

import pytest


# ── 重いモジュールをインポート前にスタブ化 ────────────────────────────
# WhisperModel / pyannote / torch など GPU依存ライブラリを差し替える
def _make_stub_module(name: str, **attrs):
    mod = types.ModuleType(name)
    for k, v in attrs.items():
        setattr(mod, k, v)
    return mod


# faster_whisper
_fw = _make_stub_module("faster_whisper")
_fw.WhisperModel = MagicMock()
sys.modules.setdefault("faster_whisper", _fw)

# torch / torchaudio / soundfile（pyannoteが要求する）
for _name in ["torch", "torchaudio", "soundfile", "pyannote", "pyannote.audio",
              "huggingface_hub"]:
    sys.modules.setdefault(_name, _make_stub_module(_name))

# アプリ本体をインポート（上記スタブが差し込まれた状態で）
with patch("faster_whisper.WhisperModel"):
    import app as _app  # noqa: E402


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 1. _assign_speaker_labels
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

class TestAssignSpeakerLabels:
    def test_basic_assignment(self):
        whisper = [{"start": 0.0, "end": 2.0, "text": "hello"}]
        diarize = [{"start": 0.0, "end": 3.0, "speaker": "SPEAKER_00"}]
        result = _app._assign_speaker_labels(whisper, diarize)
        assert result[0]["speaker"] == "SPEAKER_00"

    def test_selects_best_overlap(self):
        """2つの話者セグメントのうち重なりが大きい方を選ぶ"""
        whisper = [{"start": 1.0, "end": 3.0, "text": "hi"}]
        diarize = [
            {"start": 0.0, "end": 1.5, "speaker": "SPEAKER_00"},  # overlap=0.5
            {"start": 1.5, "end": 4.0, "speaker": "SPEAKER_01"},  # overlap=1.5
        ]
        result = _app._assign_speaker_labels(whisper, diarize)
        assert result[0]["speaker"] == "SPEAKER_01"

    def test_empty_diarize_returns_fallback(self):
        """話者セグメントが空のとき話者?になる"""
        whisper = [{"start": 0.0, "end": 1.0, "text": "test"}]
        result = _app._assign_speaker_labels(whisper, [])
        assert result[0]["speaker"] == "話者?"

    def test_empty_whisper_returns_empty(self):
        diarize = [{"start": 0.0, "end": 5.0, "speaker": "SPEAKER_00"}]
        result = _app._assign_speaker_labels([], diarize)
        assert result == []

    def test_preserves_original_fields(self):
        whisper = [{"start": 0.5, "end": 1.5, "text": "abc"}]
        diarize = [{"start": 0.0, "end": 2.0, "speaker": "SPEAKER_00"}]
        result = _app._assign_speaker_labels(whisper, diarize)
        assert result[0]["text"] == "abc"
        assert result[0]["start"] == 0.5
        assert result[0]["end"] == 1.5

    def test_multiple_segments(self):
        whisper = [
            {"start": 0.0, "end": 2.0, "text": "one"},
            {"start": 5.0, "end": 7.0, "text": "two"},
        ]
        diarize = [
            {"start": 0.0, "end": 3.0, "speaker": "SPEAKER_00"},
            {"start": 4.0, "end": 8.0, "speaker": "SPEAKER_01"},
        ]
        result = _app._assign_speaker_labels(whisper, diarize)
        assert result[0]["speaker"] == "SPEAKER_00"
        assert result[1]["speaker"] == "SPEAKER_01"


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 2. _format_speaker_transcript
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

class TestFormatSpeakerTranscript:
    def test_maps_speaker_to_numbered_label(self):
        segs = [{"start": 0.0, "end": 2.0, "text": "hello", "speaker": "SPEAKER_00"}]
        result = _app._format_speaker_transcript(segs)
        assert "話者1" in result
        assert "hello" in result

    def test_same_speaker_keeps_same_number(self):
        segs = [
            {"start": 0.0, "end": 1.0, "text": "a", "speaker": "SPEAKER_00"},
            {"start": 1.0, "end": 2.0, "text": "b", "speaker": "SPEAKER_00"},
        ]
        result = _app._format_speaker_transcript(segs)
        assert result.count("話者1") == 2
        assert "話者2" not in result

    def test_different_speakers_get_incrementing_numbers(self):
        segs = [
            {"start": 0.0, "end": 1.0, "text": "a", "speaker": "SPEAKER_00"},
            {"start": 1.0, "end": 2.0, "text": "b", "speaker": "SPEAKER_01"},
        ]
        result = _app._format_speaker_transcript(segs)
        assert "話者1" in result
        assert "話者2" in result

    def test_timestamp_format(self):
        segs = [{"start": 65.0, "end": 70.0, "text": "x", "speaker": "SPEAKER_00"}]
        result = _app._format_speaker_transcript(segs)
        assert "[01:05]" in result

    def test_empty_returns_empty_string(self):
        assert _app._format_speaker_transcript([]) == ""


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 3. extract_speakers
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

class TestExtractSpeakers:
    def test_basic_extraction(self):
        transcript = "[00:00] 話者1: こんにちは\n[00:05] 話者2: よろしく"
        result = _app.extract_speakers(transcript)
        assert "話者1" in result
        assert "話者2" in result

    def test_colon_not_included_in_label(self):
        """コロンが話者ラベルに含まれないことを確認"""
        transcript = "[00:00] 話者1: hello"
        result = _app.extract_speakers(transcript)
        assert "話者1:" not in result
        assert "話者1" in result

    def test_deduplication(self):
        transcript = "[00:00] 話者1: a\n[00:05] 話者1: b"
        result = _app.extract_speakers(transcript)
        assert result.count("話者1") == 1

    def test_preserves_insertion_order(self):
        transcript = "[00:00] 話者1: a\n[00:05] 話者2: b\n[00:10] 話者1: c"
        result = _app.extract_speakers(transcript)
        assert result == ["話者1", "話者2"]

    def test_empty_transcript(self):
        assert _app.extract_speakers("") == []

    def test_no_speakers(self):
        assert _app.extract_speakers("[00:00] こんにちは") == []


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 4. apply_speaker_names
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

class TestApplySpeakerNames:
    def test_basic_replacement(self):
        result = _app.apply_speaker_names("話者1: hello", {"話者1": "田中"})
        assert "田中: hello" == result

    def test_empty_name_skipped(self):
        """値が空文字のエントリは置換しない"""
        result = _app.apply_speaker_names("話者1: hello", {"話者1": ""})
        assert "話者1: hello" == result

    def test_longer_label_replaced_first(self):
        """話者1と話者10が混在するとき、話者10が先に処理されてサブストリング衝突しない"""
        text = "話者1: a\n話者10: b"
        result = _app.apply_speaker_names(text, {"話者1": "田中", "話者10": "鈴木"})
        assert "田中: a" in result
        assert "鈴木: b" in result
        # 「田中0: b」のような誤置換がないこと
        assert "田中0" not in result

    def test_whitespace_stripped_from_name(self):
        result = _app.apply_speaker_names("話者1: hi", {"話者1": "  田中  "})
        assert "田中: hi" == result

    def test_no_map_returns_unchanged(self):
        text = "話者1: hello"
        assert _app.apply_speaker_names(text, {}) == text


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 5. _split_transcript
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

class TestSplitTranscript:
    def test_short_text_returns_single_chunk(self):
        text = "line1\nline2\nline3"
        result = _app._split_transcript(text, 1000)
        assert result == [text]

    def test_splits_at_max_chars(self):
        lines = ["x" * 50] * 10  # 合計500文字
        text = "\n".join(lines)
        result = _app._split_transcript(text, 150)
        assert len(result) > 1
        for chunk in result:
            assert len(chunk) <= 300  # 多少超えてもよいがチャンクが分割されている

    def test_empty_text_returns_empty_list(self):
        assert _app._split_transcript("", 100) == []

    def test_single_long_line_still_added(self):
        """1行がmax_charsを超えていてもcurrentが空なら追加される"""
        long_line = "a" * 200
        result = _app._split_transcript(long_line, 100)
        assert len(result) == 1
        assert result[0] == long_line

    def test_all_chunks_non_empty(self):
        text = "\n".join([f"line{i}" * 10 for i in range(50)])
        for chunk in _app._split_transcript(text, 100):
            assert chunk.strip() != ""


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 6. _extract_text
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

class TestExtractText:
    def _make_block(self, type_: str, text: str = ""):
        b = MagicMock()
        b.type = type_
        b.text = text
        return b

    def test_extracts_text_blocks(self):
        resp = MagicMock()
        resp.content = [self._make_block("text", "hello "), self._make_block("text", "world")]
        assert _app._extract_text(resp) == "hello world"

    def test_ignores_non_text_blocks(self):
        resp = MagicMock()
        resp.content = [
            self._make_block("thinking", "internal"),
            self._make_block("text", "visible"),
        ]
        assert _app._extract_text(resp) == "visible"

    def test_empty_content_returns_empty_string(self):
        resp = MagicMock()
        resp.content = []
        assert _app._extract_text(resp) == ""


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 7. auto_identify_speakers（Claude API をモック）
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

class TestAutoIdentifySpeakers:
    def _mock_response(self, text: str):
        block = MagicMock()
        block.type = "text"
        block.text = text
        resp = MagicMock()
        resp.content = [block]
        return resp

    def test_returns_speaker_map(self):
        with patch("app.ANTHROPIC_API_KEY", "dummy"), \
             patch("app.anthropic.Anthropic") as mock_client:
            mock_client.return_value.messages.create.return_value = self._mock_response(
                '{"話者1": "田中", "話者2": "佐藤"}'
            )
            result = _app.auto_identify_speakers("話者1: 田中です\n話者2: 佐藤です")
            assert result == {"話者1": "田中", "話者2": "佐藤"}

    def test_returns_empty_when_no_api_key(self):
        with patch("app.ANTHROPIC_API_KEY", ""):
            assert _app.auto_identify_speakers("some text") == {}

    def test_returns_empty_on_malformed_json(self):
        with patch("app.ANTHROPIC_API_KEY", "dummy"), \
             patch("app.anthropic.Anthropic") as mock_client:
            mock_client.return_value.messages.create.return_value = self._mock_response(
                "JSONではないテキスト"
            )
            result = _app.auto_identify_speakers("話者1: hello")
            assert result == {}

    def test_extracts_json_from_mixed_response(self):
        """JSON前後に余分なテキストがあっても抽出できる"""
        with patch("app.ANTHROPIC_API_KEY", "dummy"), \
             patch("app.anthropic.Anthropic") as mock_client:
            mock_client.return_value.messages.create.return_value = self._mock_response(
                'こちらが結果です:\n{"話者1": "田中"}\n以上です。'
            )
            result = _app.auto_identify_speakers("話者1: hello")
            assert result == {"話者1": "田中"}


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 8. transcribe_only
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

class TestTranscribeOnly:
    def test_formats_timestamps_correctly(self):
        segments = [
            {"start": 0.0, "end": 2.0, "text": "hello"},
            {"start": 65.0, "end": 67.0, "text": "world"},
        ]
        with patch("app.transcribe_with_segments", return_value=(segments, "ja")):
            text, lang = _app.transcribe_only(Path("dummy.wav"))
        assert "[00:00] hello" in text
        assert "[01:05] world" in text
        assert lang == "ja"

    def test_empty_segments_returns_empty_string(self):
        with patch("app.transcribe_with_segments", return_value=([], "ja")):
            text, lang = _app.transcribe_only(Path("dummy.wav"))
        assert text == ""


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 9. APIエンドポイント（FastAPI TestClient）
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

try:
    from fastapi.testclient import TestClient
    import io

    _client = TestClient(_app.app, raise_server_exceptions=False)

    class TestDiarizationAvailable:
        def test_returns_false_when_no_hf_token(self):
            with patch("app.HF_TOKEN", ""):
                resp = _client.get("/diarization-available")
            assert resp.status_code == 200
            assert resp.json()["available"] is False

        def test_returns_true_when_hf_token_set(self):
            with patch("app.HF_TOKEN", "dummy_token"):
                resp = _client.get("/diarization-available")
            assert resp.status_code == 200
            assert resp.json()["available"] is True

    class TestDownloadEndpoints:
        def test_download_md_not_found(self):
            resp = _client.get("/download/md/nonexistent_file")
            assert resp.status_code == 404

        def test_download_docx_not_found(self):
            resp = _client.get("/download/docx/nonexistent_file")
            assert resp.status_code == 404

    class TestIdentifySpeakers:
        def test_returns_400_when_no_api_key(self):
            with patch("app.ANTHROPIC_API_KEY", ""):
                resp = _client.post(
                    "/identify-speakers",
                    json={"transcript": "話者1: hello"},
                )
            assert resp.status_code == 400

    class TestTranscribeStart:
        def test_rejects_unsupported_extension(self):
            resp = _client.post(
                "/transcribe-start",
                files={"file": ("test.txt", b"content", "text/plain")},
                data={"use_diarization": "false"},
            )
            assert resp.status_code == 400

        def test_accepts_mp3_and_returns_job_id(self):
            with patch("app._run_job"), \
                 patch("app.asyncio.create_task"):
                resp = _client.post(
                    "/transcribe-start",
                    files={"file": ("test.mp3", b"fake_audio", "audio/mpeg")},
                    data={"use_diarization": "false"},
                )
            assert resp.status_code == 200
            assert "job_id" in resp.json()

    class TestProgressEndpoint:
        def test_unknown_job_id_returns_error(self):
            # SSEストリームの最初のメッセージにエラーが含まれることを確認
            with _client.stream("GET", "/progress/nonexistent-job-id") as resp:
                assert resp.status_code == 200
                first_line = next(
                    (line for line in resp.iter_lines() if line.startswith("data:")),
                    None,
                )
            assert first_line is not None
            data = json.loads(first_line.removeprefix("data: "))
            assert data.get("done") is True
            assert "error" in data

    class TestIndexPage:
        def test_returns_html(self):
            resp = _client.get("/")
            assert resp.status_code == 200
            assert "text/html" in resp.headers["content-type"]
            assert "議事録自動生成" in resp.text

except ImportError:
    pass  # httpx / fastapi.testclient が未インストールの場合はスキップ


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 10. create_docx
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

class TestCreateDocx:
    """create_docx は outputs/ にファイルを書く。各テスト後に削除する。"""

    def _run(self, md: str, title: str = "test_output") -> Path:
        path = _app.create_docx(md, title)
        return path

    def _cleanup(self, path: Path):
        if path.exists():
            path.unlink()

    def test_creates_file(self):
        path = self._run("# タイトル", "test_creates_file")
        try:
            assert path.exists()
            assert path.suffix == ".docx"
        finally:
            self._cleanup(path)

    def test_headings_are_written(self):
        from docx import Document as _Document
        md = "# 見出し1\n## 見出し2\n### 見出し3"
        path = self._run(md, "test_headings")
        try:
            doc = _Document(str(path))
            styles = [p.style.name for p in doc.paragraphs]
            assert any("Heading 1" in s for s in styles)
            assert any("Heading 2" in s for s in styles)
            assert any("Heading 3" in s for s in styles)
        finally:
            self._cleanup(path)

    def test_bullet_list_written(self):
        from docx import Document as _Document
        md = "- アイテム1\n- アイテム2"
        path = self._run(md, "test_bullets")
        try:
            doc = _Document(str(path))
            texts = [p.text for p in doc.paragraphs]
            assert "アイテム1" in texts
            assert "アイテム2" in texts
        finally:
            self._cleanup(path)

    def test_table_written(self):
        from docx import Document as _Document
        md = "| 担当者 | タスク | 期限 |\n|--------|--------|------|\n| 田中 | 資料作成 | 来週 |"
        path = self._run(md, "test_table")
        try:
            doc = _Document(str(path))
            assert len(doc.tables) == 1
            row = doc.tables[0].rows[0]
            assert row.cells[0].text == "担当者"
            assert row.cells[1].text == "タスク"
        finally:
            self._cleanup(path)

    def test_separator_row_skipped(self):
        """Markdownのテーブル区切り行（---|---）はdocxに追加されない"""
        from docx import Document as _Document
        md = "| A | B |\n|---|---|\n| 1 | 2 |"
        path = self._run(md, "test_separator")
        try:
            doc = _Document(str(path))
            assert len(doc.tables) == 1
            assert len(doc.tables[0].rows) == 2  # ヘッダー行 + データ行のみ
        finally:
            self._cleanup(path)

    def test_horizontal_rule_written(self):
        from docx import Document as _Document
        md = "---"
        path = self._run(md, "test_hr")
        try:
            doc = _Document(str(path))
            texts = [p.text for p in doc.paragraphs]
            assert any("─" in t for t in texts)
        finally:
            self._cleanup(path)

    def test_plain_paragraph_written(self):
        from docx import Document as _Document
        md = "これは通常の段落です。"
        path = self._run(md, "test_plain")
        try:
            doc = _Document(str(path))
            texts = [p.text for p in doc.paragraphs]
            assert "これは通常の段落です。" in texts
        finally:
            self._cleanup(path)

    def test_empty_lines_skipped(self):
        """空行はdocxに追加されない"""
        from docx import Document as _Document
        md = "段落1\n\n\n段落2"
        path = self._run(md, "test_empty_lines")
        try:
            doc = _Document(str(path))
            texts = [p.text for p in doc.paragraphs if p.text.strip()]
            assert texts == ["段落1", "段落2"]
        finally:
            self._cleanup(path)


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 11. generate_minutes
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

class TestGenerateMinutes:
    def _mock_text_response(self, text: str):
        block = MagicMock()
        block.type = "text"
        block.text = text
        resp = MagicMock()
        resp.content = [block]
        return resp

    def _mock_stream(self, texts: list[str]):
        stream = MagicMock()
        stream.__enter__ = MagicMock(return_value=stream)
        stream.__exit__ = MagicMock(return_value=False)
        stream.text_stream = iter(texts)
        return stream

    def test_raises_without_api_key(self):
        with patch("app.ANTHROPIC_API_KEY", ""):
            with pytest.raises(RuntimeError, match="ANTHROPIC_API_KEY"):
                _app.generate_minutes("テスト")

    def test_short_transcript_no_callback(self):
        """短いテキスト・progress_cbなし → messages.create を1回呼ぶ"""
        with patch("app.ANTHROPIC_API_KEY", "dummy"), \
             patch("app.anthropic.Anthropic") as mock_cls:
            mock_cls.return_value.messages.create.return_value = self._mock_text_response("議事録テスト")
            result = _app.generate_minutes("短い文字起こし")
        assert result == "議事録テスト"
        mock_cls.return_value.messages.create.assert_called_once()

    def test_short_transcript_with_callback(self):
        """短いテキスト・progress_cbあり → messages.stream を使う"""
        with patch("app.ANTHROPIC_API_KEY", "dummy"), \
             patch("app.anthropic.Anthropic") as mock_cls:
            mock_cls.return_value.messages.stream.return_value = self._mock_stream(["議事", "録"])
            calls = []
            _app.generate_minutes("短い文字起こし", progress_cb=lambda p, l: calls.append(p))
        mock_cls.return_value.messages.stream.assert_called_once()
        assert len(calls) >= 1
        assert calls[0] == 75  # 最初のprogress_cb呼び出し

    def test_template_included_in_system_prompt(self):
        """templateが渡された場合、system_promptにテンプレートが含まれる"""
        with patch("app.ANTHROPIC_API_KEY", "dummy"), \
             patch("app.anthropic.Anthropic") as mock_cls:
            mock_cls.return_value.messages.create.return_value = self._mock_text_response("result")
            _app.generate_minutes("文字起こし", template="## カスタムフォーマット")
        call_kwargs = mock_cls.return_value.messages.create.call_args[1]
        assert "カスタムフォーマット" in call_kwargs["system"]

    def test_long_transcript_splits_into_chunks(self):
        """80,000文字超のテキストはチャンク分割して複数回APIを呼ぶ"""
        long_transcript = "a\n" * 50_000  # 約100,000文字
        with patch("app.ANTHROPIC_API_KEY", "dummy"), \
             patch("app.anthropic.Anthropic") as mock_cls:
            mock_cls.return_value.messages.create.return_value = self._mock_text_response("chunk summary")
            _app.generate_minutes(long_transcript)
        # チャンク数+最終生成で複数回呼ばれることを確認
        assert mock_cls.return_value.messages.create.call_count >= 2

    def test_long_transcript_progress_callback_called(self):
        """長いテキスト・progress_cbあり → チャンクごとにcbが呼ばれる"""
        long_transcript = "a\n" * 50_000
        with patch("app.ANTHROPIC_API_KEY", "dummy"), \
             patch("app.anthropic.Anthropic") as mock_cls:
            mock_cls.return_value.messages.create.return_value = self._mock_text_response("summary")
            mock_cls.return_value.messages.stream.return_value = self._mock_stream(["最終議事録"])
            calls = []
            _app.generate_minutes(long_transcript, progress_cb=lambda p, l: calls.append((p, l)))
        assert any("要約" in label for _, label in calls)
        assert any("最終" in label for _, label in calls)


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 12. transcribe_with_segments
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

class TestTranscribeWithSegments:
    def _seg(self, start, end, text):
        s = MagicMock()
        s.start, s.end, s.text = start, end, f"  {text}  "
        return s

    def _info(self, duration, language):
        i = MagicMock()
        i.duration, i.language = duration, language
        return i

    def test_returns_segments_and_language(self):
        _app.whisper_model.transcribe.return_value = (
            [self._seg(0.0, 2.0, "hello")], self._info(10.0, "ja")
        )
        result, lang = _app.transcribe_with_segments(Path("dummy.wav"))
        assert lang == "ja"
        assert result == [{"start": 0.0, "end": 2.0, "text": "hello"}]

    def test_whitespace_stripped_from_text(self):
        _app.whisper_model.transcribe.return_value = (
            [self._seg(0.0, 1.0, "  trimmed  ")], self._info(10.0, "ja")
        )
        result, _ = _app.transcribe_with_segments(Path("dummy.wav"))
        assert result[0]["text"] == "trimmed"

    def test_progress_pct_calculation(self):
        # end=100, total=200 → pct = 10 + int(100/200 * 50) = 35
        _app.whisper_model.transcribe.return_value = (
            [self._seg(0.0, 100.0, "x")], self._info(200.0, "ja")
        )
        calls = []
        _app.transcribe_with_segments(Path("dummy.wav"), progress_cb=lambda p, l: calls.append(p))
        assert calls[0] == min(58, 10 + int(100.0 / 200.0 * 50))

    def test_pct_capped_at_58(self):
        # end >> total → without cap would exceed 58
        _app.whisper_model.transcribe.return_value = (
            [self._seg(0.0, 999.0, "x")], self._info(1.0, "ja")
        )
        calls = []
        _app.transcribe_with_segments(Path("dummy.wav"), progress_cb=lambda p, l: calls.append(p))
        assert calls[0] == 58

    def test_progress_label_contains_timestamps(self):
        _app.whisper_model.transcribe.return_value = (
            [self._seg(0.0, 65.0, "x")], self._info(180.0, "ja")
        )
        labels = []
        _app.transcribe_with_segments(Path("dummy.wav"), progress_cb=lambda p, l: labels.append(l))
        assert "01:05" in labels[0]
        assert "03:00" in labels[0]

    def test_zero_duration_does_not_raise(self):
        # duration=0 → max(0 or 1, 1) = 1, no ZeroDivisionError
        _app.whisper_model.transcribe.return_value = (
            [self._seg(0.0, 0.0, "x")], self._info(0, "ja")
        )
        result, _ = _app.transcribe_with_segments(Path("dummy.wav"))
        assert len(result) == 1

    def test_no_progress_cb_works(self):
        _app.whisper_model.transcribe.return_value = (
            [self._seg(0.0, 1.0, "hi")], self._info(5.0, "en")
        )
        result, lang = _app.transcribe_with_segments(Path("dummy.wav"))
        assert lang == "en"


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 13. generate_speaker_summaries
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

class TestGenerateSpeakerSummaries:
    def _mock_response(self, text):
        block = MagicMock()
        block.type, block.text = "text", text
        resp = MagicMock()
        resp.content = [block]
        return resp

    def test_raises_without_api_key(self):
        with patch("app.ANTHROPIC_API_KEY", ""):
            with pytest.raises(RuntimeError, match="ANTHROPIC_API_KEY"):
                _app.generate_speaker_summaries("transcript")

    def test_returns_summary_text(self):
        with patch("app.ANTHROPIC_API_KEY", "dummy"), \
             patch("app.anthropic.Anthropic") as mock_cls:
            mock_cls.return_value.messages.create.return_value = self._mock_response("## 話者別サマリー\n内容")
            result = _app.generate_speaker_summaries("話者1: こんにちは")
        assert "話者別サマリー" in result

    def test_uses_speaker_summary_system_prompt(self):
        with patch("app.ANTHROPIC_API_KEY", "dummy"), \
             patch("app.anthropic.Anthropic") as mock_cls:
            mock_cls.return_value.messages.create.return_value = self._mock_response("result")
            _app.generate_speaker_summaries("transcript")
        call_kwargs = mock_cls.return_value.messages.create.call_args[1]
        assert call_kwargs["system"] == _app.SPEAKER_SUMMARY_SYSTEM

    def test_passes_transcript_in_user_message(self):
        with patch("app.ANTHROPIC_API_KEY", "dummy"), \
             patch("app.anthropic.Anthropic") as mock_cls:
            mock_cls.return_value.messages.create.return_value = self._mock_response("ok")
            _app.generate_speaker_summaries("話者1: テスト発言")
        messages = mock_cls.return_value.messages.create.call_args[1]["messages"]
        assert "話者1: テスト発言" in messages[0]["content"]


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 14. diarize
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

class TestDiarize:
    def _make_pipeline(self, segments):
        mock_diarization = MagicMock()
        mock_diarization.itertracks.return_value = segments
        return MagicMock(return_value=mock_diarization)

    def _turn(self, start, end):
        t = MagicMock()
        t.start, t.end = start, end
        return t

    def test_wav_skips_ffmpeg_conversion(self):
        pipeline = self._make_pipeline([(self._turn(0.0, 5.0), None, "SPEAKER_00")])
        with patch("app.get_diarization_pipeline", return_value=pipeline), \
             patch("app.subprocess.run") as mock_run:
            result = _app.diarize(Path("audio.wav"))
        mock_run.assert_not_called()
        assert result == [{"start": 0.0, "end": 5.0, "speaker": "SPEAKER_00"}]

    def test_non_wav_calls_ffmpeg(self):
        pipeline = self._make_pipeline([(self._turn(0.0, 3.0), None, "SPEAKER_00")])
        with patch("app.get_diarization_pipeline", return_value=pipeline), \
             patch("app.subprocess.run") as mock_run:
            _app.diarize(Path("audio.mp3"))
        mock_run.assert_called_once()
        cmd = mock_run.call_args[0][0]
        assert any("audio.mp3" in str(a) for a in cmd)

    def test_returns_multiple_segments(self):
        turns = [
            (self._turn(0.0, 2.0), None, "SPEAKER_00"),
            (self._turn(2.5, 5.0), None, "SPEAKER_01"),
        ]
        pipeline = self._make_pipeline(turns)
        with patch("app.get_diarization_pipeline", return_value=pipeline), \
             patch("app.subprocess.run"):
            result = _app.diarize(Path("audio.wav"))
        assert len(result) == 2
        assert result[0] == {"start": 0.0, "end": 2.0, "speaker": "SPEAKER_00"}
        assert result[1] == {"start": 2.5, "end": 5.0, "speaker": "SPEAKER_01"}

    def test_empty_pipeline_result_returns_empty_list(self):
        pipeline = self._make_pipeline([])
        with patch("app.get_diarization_pipeline", return_value=pipeline), \
             patch("app.subprocess.run"):
            result = _app.diarize(Path("audio.wav"))
        assert result == []

    def test_wav_path_derived_from_input(self):
        """非WAVのとき、変換先パスは入力ファイルの.wav版になる"""
        pipeline = self._make_pipeline([])
        with patch("app.get_diarization_pipeline", return_value=pipeline), \
             patch("app.subprocess.run") as mock_run:
            _app.diarize(Path("audio.m4a"))
        cmd = mock_run.call_args[0][0]
        assert any("audio.wav" in str(a) for a in cmd)


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 15. POST /transcribe  16. POST /process-speakers (TestClient)
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

try:
    from fastapi.testclient import TestClient as _TC2

    _client2 = _TC2(_app.app, raise_server_exceptions=False)

    class TestTranscribeEndpoint:
        def setup_method(self):
            self.tmp_dir = Path(tempfile.mkdtemp())
            self._orig_dir = _app.OUTPUT_DIR
            _app.OUTPUT_DIR = self.tmp_dir

        def teardown_method(self):
            _app.OUTPUT_DIR = self._orig_dir
            shutil.rmtree(self.tmp_dir, ignore_errors=True)

        def test_rejects_unsupported_extension(self):
            resp = _client2.post(
                "/transcribe",
                files={"file": ("test.txt", b"content", "text/plain")},
                data={"use_diarization": "false"},
            )
            assert resp.status_code == 400

        def test_no_diarization_returns_success(self):
            with patch("app.transcribe_only", return_value=("文字起こし内容", "ja")), \
                 patch("app.generate_minutes", return_value="# 議事録"), \
                 patch("app.create_docx"):
                resp = _client2.post(
                    "/transcribe",
                    files={"file": ("audio.mp3", b"fake", "audio/mpeg")},
                    data={"use_diarization": "false"},
                )
            assert resp.status_code == 200
            data = resp.json()
            assert data["success"] is True
            assert data["diarization_used"] is False
            assert data["transcript"] == "文字起こし内容"
            assert data["language"] == "ja"

        def test_diarization_returns_success(self):
            whisper_segs = [{"start": 0.0, "end": 2.0, "text": "hello"}]
            diarize_segs = [{"start": 0.0, "end": 3.0, "speaker": "SPEAKER_00"}]
            with patch("app.transcribe_with_segments", return_value=(whisper_segs, "ja")), \
                 patch("app.diarize", return_value=diarize_segs), \
                 patch("app.generate_minutes", return_value="# 議事録"), \
                 patch("app.create_docx"):
                resp = _client2.post(
                    "/transcribe",
                    files={"file": ("audio.wav", b"fake", "audio/wav")},
                    data={"use_diarization": "true"},
                )
            assert resp.status_code == 200
            assert resp.json()["diarization_used"] is True

        def test_template_file_passed_to_generate_minutes(self):
            with patch("app.transcribe_only", return_value=("transcript", "ja")), \
                 patch("app.generate_minutes", return_value="# 議事録") as mock_gen, \
                 patch("app.create_docx"):
                resp = _client2.post(
                    "/transcribe",
                    files={
                        "file": ("audio.mp3", b"fake", "audio/mpeg"),
                        "template_file": ("tmpl.md", "## テンプレート".encode(), "text/markdown"),
                    },
                    data={"use_diarization": "false"},
                )
            assert resp.status_code == 200
            template_arg = mock_gen.call_args[0][1]
            assert template_arg != ""

        def test_response_includes_download_links(self):
            with patch("app.transcribe_only", return_value=("text", "ja")), \
                 patch("app.generate_minutes", return_value="# 議事録"), \
                 patch("app.create_docx"):
                resp = _client2.post(
                    "/transcribe",
                    files={"file": ("audio.mp3", b"fake", "audio/mpeg")},
                    data={"use_diarization": "false"},
                )
            data = resp.json()
            assert data["files"]["markdown"].startswith("/download/md/")
            assert data["files"]["word"].startswith("/download/docx/")

    class TestProcessSpeakersEndpoint:
        def setup_method(self):
            self.tmp_dir = Path(tempfile.mkdtemp())
            self._orig_dir = _app.OUTPUT_DIR
            _app.OUTPUT_DIR = self.tmp_dir

        def teardown_method(self):
            _app.OUTPUT_DIR = self._orig_dir
            shutil.rmtree(self.tmp_dir, ignore_errors=True)

        def test_replaces_speaker_names_in_transcript_and_minutes(self):
            with patch("app.generate_speaker_summaries", return_value="## サマリー"), \
                 patch("app.create_docx"):
                resp = _client2.post(
                    "/process-speakers",
                    json={
                        "transcript": "話者1: おはよう\n話者2: こんにちは",
                        "minutes": "# 議事録\n話者1が発言した",
                        "speaker_map": {"話者1": "田中", "話者2": "佐藤"},
                        "title": "test_speakers",
                    },
                )
            assert resp.status_code == 200
            data = resp.json()
            assert data["success"] is True
            assert "田中" in data["transcript"]
            assert "佐藤" in data["transcript"]
            assert "田中" in data["minutes"]

        def test_empty_speaker_map_leaves_labels_unchanged(self):
            with patch("app.generate_speaker_summaries", return_value="summary"), \
                 patch("app.create_docx"):
                resp = _client2.post(
                    "/process-speakers",
                    json={
                        "transcript": "話者1: hello",
                        "minutes": "# 議事録",
                        "speaker_map": {},
                        "title": "test_no_map",
                    },
                )
            assert resp.status_code == 200
            assert resp.json()["transcript"] == "話者1: hello"

        def test_empty_title_generates_default_prefix(self):
            with patch("app.generate_speaker_summaries", return_value="summary"), \
                 patch("app.create_docx"):
                resp = _client2.post(
                    "/process-speakers",
                    json={
                        "transcript": "話者1: hi",
                        "minutes": "# 議事録",
                        "speaker_map": {},
                    },
                )
            assert resp.status_code == 200
            assert "speakers_" in resp.json()["files"]["markdown"]

        def test_speaker_summary_included_in_response(self):
            with patch("app.generate_speaker_summaries", return_value="## 話者別サマリー\n内容"), \
                 patch("app.create_docx"):
                resp = _client2.post(
                    "/process-speakers",
                    json={
                        "transcript": "話者1: hi",
                        "minutes": "# 議事録",
                        "speaker_map": {},
                        "title": "test_summary",
                    },
                )
            assert "話者別サマリー" in resp.json()["speaker_summary"]

except ImportError:
    pass


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 17. _run_job (async background job)
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

class TestRunJob:
    pytestmark = pytest.mark.asyncio

    def setup_method(self):
        self.tmp_dir = Path(tempfile.mkdtemp())
        self._orig_dir = _app.OUTPUT_DIR
        _app.OUTPUT_DIR = self.tmp_dir

    def teardown_method(self):
        _app.OUTPUT_DIR = self._orig_dir
        shutil.rmtree(self.tmp_dir, ignore_errors=True)

    @pytest.mark.asyncio
    async def test_no_diarize_job_completes_successfully(self):
        job_id = "run-job-1"
        tmp = Path(tempfile.mktemp(suffix=".mp3"))
        tmp.write_bytes(b"fake")
        _app._jobs[job_id] = {"pct": 0, "label": "", "done": False}

        with patch("app.transcribe_with_segments",
                   return_value=([{"start": 0.0, "end": 1.0, "text": "hello"}], "ja")), \
             patch("app.generate_minutes", return_value="# 議事録"), \
             patch("app.create_docx"):
            await _app._run_job(job_id, tmp, False, "", "test.mp3")

        job = _app._jobs[job_id]
        assert job["done"] is True
        assert job["result"]["success"] is True
        assert job["result"]["language"] == "ja"
        assert job["result"]["diarization_used"] is False
        assert not tmp.exists()
        _app._jobs.pop(job_id, None)

    @pytest.mark.asyncio
    async def test_diarize_mode_calls_diarize_function(self):
        job_id = "run-job-2"
        tmp = Path(tempfile.mktemp(suffix=".wav"))
        tmp.write_bytes(b"fake")
        _app._jobs[job_id] = {"pct": 0, "label": "", "done": False}

        with patch("app.transcribe_with_segments",
                   return_value=([{"start": 0.0, "end": 1.0, "text": "hi"}], "ja")), \
             patch("app.diarize",
                   return_value=[{"start": 0.0, "end": 1.0, "speaker": "SPEAKER_00"}]) as mock_diarize, \
             patch("app.generate_minutes", return_value="# 議事録"), \
             patch("app.create_docx"):
            await _app._run_job(job_id, tmp, True, "", "test.wav")

        mock_diarize.assert_called_once_with(tmp)
        assert _app._jobs[job_id]["result"]["diarization_used"] is True
        _app._jobs.pop(job_id, None)

    @pytest.mark.asyncio
    async def test_error_sets_job_error_state(self):
        job_id = "run-job-3"
        tmp = Path(tempfile.mktemp(suffix=".mp3"))
        _app._jobs[job_id] = {"pct": 0, "label": "", "done": False}

        with patch("app.transcribe_with_segments", side_effect=RuntimeError("transcribe failed")):
            await _app._run_job(job_id, tmp, False, "", "test.mp3")

        assert _app._jobs[job_id]["done"] is True
        assert "transcribe failed" in _app._jobs[job_id]["error"]
        _app._jobs.pop(job_id, None)

    @pytest.mark.asyncio
    async def test_tmp_file_cleaned_up_on_error(self):
        job_id = "run-job-4"
        tmp = Path(tempfile.mktemp(suffix=".mp3"))
        tmp.write_bytes(b"fake")
        _app._jobs[job_id] = {"pct": 0, "label": "", "done": False}

        with patch("app.transcribe_with_segments", side_effect=RuntimeError("fail")):
            await _app._run_job(job_id, tmp, False, "", "test.mp3")

        assert not tmp.exists()
        _app._jobs.pop(job_id, None)

    @pytest.mark.asyncio
    async def test_with_template_text(self):
        job_id = "run-job-5"
        tmp = Path(tempfile.mktemp(suffix=".mp3"))
        tmp.write_bytes(b"fake")
        _app._jobs[job_id] = {"pct": 0, "label": "", "done": False}

        with patch("app.transcribe_with_segments",
                   return_value=([{"start": 0.0, "end": 1.0, "text": "hello"}], "ja")), \
             patch("app.generate_minutes", return_value="# 議事録") as mock_gen, \
             patch("app.create_docx"):
            await _app._run_job(job_id, tmp, False, "## テンプレート", "test.mp3")

        template_arg = mock_gen.call_args[0][1]
        assert template_arg == "## テンプレート"
        assert _app._jobs[job_id]["result"]["template_used"] is True
        _app._jobs.pop(job_id, None)
