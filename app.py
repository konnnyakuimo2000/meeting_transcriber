"""
meeting_transcriber/app.py
長時間音声ファイルを自動で議事録化するWebアプリ（話者識別対応）

起動: uvicorn app:app --reload --port 8510

環境変数:
  # AIプロバイダー（いずれか1つのAPIキーを設定。複数ある場合は AI_PROVIDER で指定）
  AI_PROVIDER         ... anthropic / openai / gemini（省略時は設定済みキーから自動検出）
  ANTHROPIC_API_KEY   ... Claude APIキー
  OPENAI_API_KEY      ... OpenAI APIキー
  OPENAI_MODEL        ... 使用モデル（デフォルト: gpt-4o）
  GEMINI_API_KEY      ... Google Gemini APIキー
  GEMINI_MODEL        ... 使用モデル（デフォルト: gemini-1.5-pro）
  HF_TOKEN            ... HuggingFace Access Token（話者識別に必要）
  WHISPER_MODEL       ... tiny/base/small/medium/large-v3 (デフォルト: medium)
"""
from __future__ import annotations

import os
import asyncio
import subprocess
import tempfile
import uuid
from pathlib import Path
from datetime import datetime

# torchaudio 2.11+ removed several APIs that pyannote.audio 3.x still uses.
# Each patch is independent so one failure doesn't block the rest.
try:
    import torchaudio as _torchaudio
    import soundfile as _sf
    import torch as _torch_pa

    if not hasattr(_torchaudio, "AudioMetaData"):
        from dataclasses import dataclass as _dc

        @_dc
        class _AudioMetaData:
            sample_rate: int
            num_frames: int
            num_channels: int
            bits_per_sample: int
            encoding: str

        _torchaudio.AudioMetaData = _AudioMetaData
except Exception:
    pass

try:
    import torchaudio as _torchaudio
    if not hasattr(_torchaudio, "list_audio_backends"):
        _torchaudio.list_audio_backends = lambda: ["soundfile", "sox_io"]
except Exception:
    pass

try:
    import torchaudio as _torchaudio
    import soundfile as _sf
    import torch as _torch_pa

    _orig_torchaudio_load = getattr(_torchaudio, "load", None)

    def _sf_load(path, frame_offset=0, num_frames=-1, **kwargs):
        kwargs.pop("backend", None)
        with _sf.SoundFile(str(path)) as f:
            sr = f.samplerate
            if frame_offset > 0:
                f.seek(frame_offset)
            frames = num_frames if num_frames > 0 else -1
            data = f.read(frames=frames, dtype="float32", always_2d=True)
        return _torch_pa.from_numpy(data.T), sr

    _torchaudio.load = _sf_load
except Exception:
    pass

try:
    import torchaudio as _torchaudio
    import soundfile as _sf

    def _sf_info(path, **kwargs):
        kwargs.pop("backend", None)
        with _sf.SoundFile(str(path)) as f:
            return _torchaudio.AudioMetaData(
                sample_rate=f.samplerate,
                num_frames=len(f),
                num_channels=f.channels,
                bits_per_sample=16,
                encoding="PCM_S",
            )

    _torchaudio.info = _sf_info
except Exception:
    pass

# PyTorch 2.6+ changed torch.load default to weights_only=True;
# pyannote checkpoints contain many custom classes not in the allowlist.
# Force weights_only=False at the serialization level (safe for trusted HuggingFace models).
try:
    import torch as _torch
    import torch.serialization as _torch_serial
    import functools as _functools

    _orig_serial_load = _torch_serial.load

    @_functools.wraps(_orig_serial_load)
    def _patched_serial_load(*args, **kwargs):
        kwargs["weights_only"] = False
        return _orig_serial_load(*args, **kwargs)

    _torch_serial.load = _patched_serial_load
    _torch.load = _patched_serial_load
except Exception:
    pass

# huggingface_hub 0.20+ removed use_auth_token; pyannote.audio 3.x still passes it.
try:
    import huggingface_hub as _hf_hub
    import functools as _functools

    for _fn_name in ("hf_hub_download", "snapshot_download"):
        _orig = getattr(_hf_hub, _fn_name, None)
        if _orig is None:
            continue

        @_functools.wraps(_orig)
        def _patched(*args, _orig=_orig, **kwargs):
            if "use_auth_token" in kwargs:
                token = kwargs.pop("use_auth_token")
                kwargs.setdefault("token", token)
            return _orig(*args, **kwargs)

        setattr(_hf_hub, _fn_name, _patched)
except Exception:
    pass

import json
import re

import aiofiles
import anthropic
from fastapi import FastAPI, File, UploadFile, Form, HTTPException
from fastapi.responses import HTMLResponse, JSONResponse, FileResponse, StreamingResponse
from pydantic import BaseModel
from faster_whisper import WhisperModel
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── 設定 ──────────────────────────────────────────────────────────────
from dotenv import load_dotenv
load_dotenv()

WHISPER_MODEL_SIZE = os.getenv("WHISPER_MODEL", "medium")
HF_TOKEN = os.getenv("HF_TOKEN", "")
ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY", "")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-4o")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "")
GEMINI_MODEL = os.getenv("GEMINI_MODEL", "gemini-1.5-pro")
_AI_PROVIDER_ENV = os.getenv("AI_PROVIDER", "").lower()
OUTPUT_DIR = Path("outputs")
OUTPUT_DIR.mkdir(exist_ok=True)

# ffmpegパス
FFMPEG_EXE = Path(r"C:\Users\toshi\AppData\Local\Microsoft\WinGet\Packages\Gyan.FFmpeg_Microsoft.Winget.Source_8wekyb3d8bbwe\ffmpeg-8.1.1-full_build\bin\ffmpeg.exe")

# ── FastAPI ───────────────────────────────────────────────────────────
app = FastAPI(title="議事録自動生成")

# ── Whisperモデル（起動時に一度だけロード）────────────────────────────
print(f"Whisperモデル ({WHISPER_MODEL_SIZE}) をロード中...")
whisper_model = WhisperModel(WHISPER_MODEL_SIZE, device="cuda", compute_type="float16")
print("Whisperモデル ロード完了")

# ── ジョブ進捗管理 ────────────────────────────────────────────────────
_jobs: dict[str, dict] = {}

# ── pyannote 話者識別モデル（遅延ロード）─────────────────────────────
_diarization_pipeline = None


def get_diarization_pipeline():
    global _diarization_pipeline
    if _diarization_pipeline is not None:
        return _diarization_pipeline
    if not HF_TOKEN:
        raise RuntimeError(
            "HF_TOKEN が設定されていません。"
            ".env に HuggingFace Access Token を設定してください。"
        )
    from pyannote.audio import Pipeline
    from huggingface_hub import login
    login(token=HF_TOKEN, add_to_git_credential=False)
    print("話者識別モデルをロード中...")
    _diarization_pipeline = Pipeline.from_pretrained(
        "pyannote/speaker-diarization-3.1",
    )
    # GPUが使えれば自動で使う
    try:
        import torch
        if torch.cuda.is_available():
            _diarization_pipeline = _diarization_pipeline.to(torch.device("cuda"))
            print("話者識別: GPU使用")
        else:
            print("話者識別: CPU使用（GPUより遅い）")
    except ImportError:
        pass
    print("話者識別モデル ロード完了")
    return _diarization_pipeline


# ── 話者識別 ─────────────────────────────────────────────────────────
def diarize(audio_path: Path) -> list[dict]:
    """
    音声ファイルを話者識別し、セグメントのリストを返す。
    [{"start": 0.0, "end": 5.2, "speaker": "SPEAKER_00"}, ...]
    pyannoteはWAVのみ対応のため、必要に応じてWAVに変換する。
    """
    pipeline = get_diarization_pipeline()

    # WAV以外はffmpegで直接変換
    wav_path = None
    if audio_path.suffix.lower() != ".wav":
        wav_path = audio_path.with_suffix(".wav")
        subprocess.run(
            [str(FFMPEG_EXE), "-y", "-i", str(audio_path),
             "-ar", "16000", "-ac", "1", str(wav_path)],
            check=True, capture_output=True,
        )
        target_path = wav_path
    else:
        target_path = audio_path

    try:
        diarization = pipeline(str(target_path))
    finally:
        if wav_path and wav_path.exists():
            wav_path.unlink()

    segments = []
    for turn, _, speaker in diarization.itertracks(yield_label=True):
        segments.append({
            "start": turn.start,
            "end": turn.end,
            "speaker": speaker,
        })
    return segments


# ── Whisper 文字起こし（セグメント付き）──────────────────────────────
def transcribe_with_segments(audio_path: Path, progress_cb=None) -> tuple[list[dict], str]:
    """
    Whisperで文字起こしし、タイムスタンプ付きセグメントのリストを返す。
    progress_cb(pct, label) が渡された場合、セグメントごとに呼び出す。
    """
    segments_gen, info = whisper_model.transcribe(
        str(audio_path),
        beam_size=5,
        language=None,
        vad_filter=True,
        vad_parameters={"min_silence_duration_ms": 500},
    )
    total = max(getattr(info, "duration", 1) or 1, 1)
    result = []
    for seg in segments_gen:
        result.append({"start": seg.start, "end": seg.end, "text": seg.text.strip()})
        if progress_cb:
            pct = min(58, 10 + int(seg.end / total * 50))
            mm_c, ss_c = int(seg.end // 60), int(seg.end % 60)
            mm_t, ss_t = int(total // 60), int(total % 60)
            progress_cb(pct, f"文字起こし中... {mm_c:02d}:{ss_c:02d} / {mm_t:02d}:{ss_t:02d}")
    return result, info.language


def transcribe_only(audio_path: Path) -> tuple[str, str]:
    """話者識別なしの文字起こし（従来モード）。"""
    segments, lang = transcribe_with_segments(audio_path)
    lines = []
    for seg in segments:
        ts = f"[{int(seg['start']//60):02d}:{int(seg['start']%60):02d}]"
        lines.append(f"{ts} {seg['text']}")
    return "\n".join(lines), lang


# ── 話者ラベルのマッピング ────────────────────────────────────────────
def _assign_speaker_labels(
    whisper_segs: list[dict],
    diarize_segs: list[dict],
) -> list[dict]:
    """
    Whisperセグメントに話者ラベルを付与する。
    各Whisperセグメントの中心時刻が最も重なるdiarizeセグメントの話者を採用。
    """
    labeled = []
    for wseg in whisper_segs:
        mid = (wseg["start"] + wseg["end"]) / 2
        best_speaker = "話者?"
        best_overlap = -1
        for dseg in diarize_segs:
            overlap = min(wseg["end"], dseg["end"]) - max(wseg["start"], dseg["start"])
            if overlap > best_overlap:
                best_overlap = overlap
                best_speaker = dseg["speaker"]
        labeled.append({**wseg, "speaker": best_speaker})
    return labeled


def _format_speaker_transcript(labeled_segs: list[dict]) -> str:
    """話者付き文字起こしをテキスト形式にフォーマット。"""
    # SPEAKER_00 → 話者1 のように番号化
    speaker_map: dict[str, str] = {}
    counter = 1
    lines = []
    for seg in labeled_segs:
        raw = seg["speaker"]
        if raw not in speaker_map:
            speaker_map[raw] = f"話者{counter}"
            counter += 1
        label = speaker_map[raw]
        ts = f"[{int(seg['start']//60):02d}:{int(seg['start']%60):02d}]"
        lines.append(f"{ts} {label}: {seg['text']}")
    return "\n".join(lines)


# ── Claude 議事録生成 ─────────────────────────────────────────────────
MINUTES_SYSTEM_PROMPT = """あなたは優秀な議事録作成アシスタントです。
音声の文字起こしテキスト（話者ラベル付き）をもとに、構造化された議事録を作成してください。

議事録フォーマット:
# 議事録

## 基本情報
- 日時: （文字起こしから読み取れる場合）
- 参加者: 話者1、話者2 ... （文字起こしに登場した話者）
- 会議名/議題: （推定）

## 議題・討議内容
（主要な議題ごとに整理。誰が何を言ったか明確に記録）

### [議題1]
- 話者1: 〜
- 話者2: 〜

## 決定事項
- （箇条書きで明確に。誰が決定したか記載）

## 課題・懸念事項
- （解決されなかった問題、今後の検討事項）

## アクションアイテム
| 担当者 | タスク | 期限 |
|--------|--------|------|
| ...    | ...    | ...  |

## 次回予定
- （次回会議の日程・議題など、言及があれば）

---
各話者の発言を正確に記録し、誰が何を言ったか明確にしてください。
不明瞭な箇所は「（不明）」と記載してください。"""


def generate_minutes(transcript: str, template: str = "", progress_cb=None) -> str:
    """AIで議事録を生成する。progress_cbが渡された場合はストリーミングで進捗を更新する。"""
    if not _has_api_key():
        raise RuntimeError("AIプロバイダーのAPIキーが設定されていません。")
    label = _provider_label()

    if template:
        system_prompt = (
            MINUTES_SYSTEM_PROMPT
            + f"\n\n## 出力フォーマット（必ずこのテンプレートに従うこと）\n\n{template}"
        )
        user_prefix = "以下のテンプレートフォーマットに厳密に従い、文字起こしから議事録を作成してください:"
    else:
        system_prompt = MINUTES_SYSTEM_PROMPT
        user_prefix = "以下の文字起こしから議事録を作成してください:"

    MAX_CHARS = 80_000
    if len(transcript) > MAX_CHARS:
        chunks = _split_transcript(transcript, MAX_CHARS)
        summaries = []
        for i, chunk in enumerate(chunks, 1):
            if progress_cb:
                progress_cb(75 + int(i / len(chunks) * 10), f"議事録生成中（第{i}/{len(chunks)}部を要約）...")
            summaries.append(_call_ai(
                system=f"あなたは議事録作成アシスタントです。長い会議の第{i}/{len(chunks)}部を要約してください。誰が何を言ったか、決定事項・アクションアイテムを含めること。",
                user=f"以下の文字起こし（第{i}部）を要約:\n\n{chunk}",
                max_tokens=4096,
            ))
        combined = "\n\n---\n\n".join(summaries)
        final_prompt = f"以下は各パートの要約です。{'テンプレートフォーマットに従い、' if template else ''}最終的な議事録を作成してください:\n\n{combined}"
        if progress_cb:
            progress_cb(87, f"最終議事録を生成中（{label}）...")
            minutes_text = ""
            for text in _call_ai_stream(system_prompt, final_prompt, max_tokens=8192):
                minutes_text += text
                progress_cb(min(94, 87 + len(minutes_text) // 200), f"議事録を生成中（{label}）...")
            return minutes_text
        else:
            return _call_ai(system_prompt, final_prompt, max_tokens=8192)
    else:
        if progress_cb:
            progress_cb(75, f"議事録を生成中（{label}）...")
            minutes_text = ""
            for text in _call_ai_stream(system_prompt, f"{user_prefix}\n\n{transcript}", max_tokens=8192):
                minutes_text += text
                progress_cb(min(94, 75 + len(minutes_text) // 150), f"議事録を生成中（{label}）...")
            return minutes_text
        else:
            return _call_ai(system_prompt, f"{user_prefix}\n\n{transcript}", max_tokens=8192)


def _extract_text(response) -> str:
    return "".join(b.text for b in response.content if b.type == "text")


# ── AIプロバイダー抽象化 ──────────────────────────────────────────────────────
def _get_provider() -> str:
    if _AI_PROVIDER_ENV in ("anthropic", "openai", "gemini"):
        return _AI_PROVIDER_ENV
    if ANTHROPIC_API_KEY:
        return "anthropic"
    if OPENAI_API_KEY:
        return "openai"
    if GEMINI_API_KEY:
        return "gemini"
    return "anthropic"


def _has_api_key() -> bool:
    p = _get_provider()
    return bool(
        (p == "anthropic" and ANTHROPIC_API_KEY)
        or (p == "openai" and OPENAI_API_KEY)
        or (p == "gemini" and GEMINI_API_KEY)
    )


def _provider_label() -> str:
    return {"anthropic": "Claude", "openai": "OpenAI", "gemini": "Gemini"}.get(_get_provider(), "AI")


def _call_ai(system: str, user: str, max_tokens: int = 4096) -> str:
    p = _get_provider()
    if p == "anthropic":
        client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
        resp = client.messages.create(
            model="claude-opus-4-8",
            max_tokens=max_tokens,
            thinking={"type": "adaptive"},
            system=system,
            messages=[{"role": "user", "content": user}],
        )
        return _extract_text(resp)
    if p == "openai":
        import openai as _openai
        client = _openai.OpenAI(api_key=OPENAI_API_KEY)
        resp = client.chat.completions.create(
            model=OPENAI_MODEL,
            max_tokens=max_tokens,
            messages=[{"role": "system", "content": system}, {"role": "user", "content": user}],
        )
        return resp.choices[0].message.content or ""
    if p == "gemini":
        import google.generativeai as genai
        genai.configure(api_key=GEMINI_API_KEY)
        m = genai.GenerativeModel(GEMINI_MODEL, system_instruction=system)
        return m.generate_content(user).text
    raise RuntimeError(f"Unknown AI provider: {p}")


def _call_ai_stream(system: str, user: str, max_tokens: int = 8192):
    """テキストチャンクをyieldするジェネレータ。"""
    p = _get_provider()
    if p == "anthropic":
        client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
        with client.messages.stream(
            model="claude-opus-4-8",
            max_tokens=max_tokens,
            thinking={"type": "adaptive"},
            system=system,
            messages=[{"role": "user", "content": user}],
        ) as stream:
            yield from stream.text_stream
    elif p == "openai":
        import openai as _openai
        client = _openai.OpenAI(api_key=OPENAI_API_KEY)
        stream = client.chat.completions.create(
            model=OPENAI_MODEL,
            max_tokens=max_tokens,
            messages=[{"role": "system", "content": system}, {"role": "user", "content": user}],
            stream=True,
        )
        for chunk in stream:
            delta = chunk.choices[0].delta.content if chunk.choices else None
            if delta:
                yield delta
    elif p == "gemini":
        import google.generativeai as genai
        genai.configure(api_key=GEMINI_API_KEY)
        m = genai.GenerativeModel(GEMINI_MODEL, system_instruction=system)
        for chunk in m.generate_content(user, stream=True):
            if chunk.text:
                yield chunk.text


# ── 話者名置換 & 話者別サマリー ───────────────────────────────────────
def extract_speakers(transcript: str) -> list[str]:
    """文字起こしテキストから話者ラベルを順番に抽出（重複なし）。"""
    seen: dict[str, None] = {}
    for m in re.finditer(r"話者[^\s:：]+", transcript):
        seen[m.group()] = None
    return list(seen.keys())


def apply_speaker_names(text: str, speaker_map: dict[str, str]) -> str:
    """テキスト中の話者ラベルを登録名に一括置換する。"""
    # 長い名前から置換してサブストリング衝突を避ける
    for label in sorted(speaker_map.keys(), key=len, reverse=True):
        name = speaker_map[label].strip()
        if name:
            text = text.replace(label, name)
    return text


def auto_identify_speakers(transcript: str) -> dict[str, str]:
    """
    文字起こしの文脈からAIが話者名を推定する。
    {"話者1": "田中", "話者2": "佐藤"} の形式で返す。
    推定できない話者は含めない。
    """
    if not _has_api_key():
        return {}
    raw = _call_ai(
        system="""あなたは会議の文字起こしから話者の実名を特定するアシスタントです。

以下のルールに従ってください：
- 自己紹介（「私は田中です」「〇〇と申します」）
- 呼びかけ（「田中さん」「〇〇くん」）
- 敬称・役職の言及（「部長の〇〇が」）
などから話者の実名を特定してください。

必ずJSON形式のみで返してください。例：
{"話者1": "田中", "話者2": "佐藤", "話者3": "鈴木"}

推定できない話者はJSONに含めないでください。
名前が全く特定できない場合は {} を返してください。
JSON以外の文章は一切出力しないでください。""",
        user=f"以下の文字起こしから話者の実名を特定してください:\n\n{transcript[:8000]}",
        max_tokens=512,
    ).strip()
    try:
        m = re.search(r"\{.*\}", raw, re.DOTALL)
        return json.loads(m.group()) if m else {}
    except Exception:
        return {}


SPEAKER_SUMMARY_SYSTEM = """あなたは会議分析アシスタントです。
話者識別済みの文字起こしをもとに、以下を出力してください。

## 話者別発言サマリー

各話者について:
1. **主な主張・意見**: その人が主に述べたこと（箇条書き3〜5点）
2. **キーワード**: 繰り返し言及したテーマや用語（3〜5個）
3. **発言スタンス**: 積極的・慎重・中立など一言で

出力はMarkdown形式で、話者ごとにセクションを分けてください。
話者名が「田中」「佐藤」などの実名の場合はそのまま使用し、
「話者1」などのラベルのままの場合もそのまま使用してください。"""


def generate_speaker_summaries(transcript: str) -> str:
    """AIで話者別サマリーを生成する。"""
    if not _has_api_key():
        raise RuntimeError("AIプロバイダーのAPIキーが設定されていません。")
    return _call_ai(
        system=SPEAKER_SUMMARY_SYSTEM,
        user=f"以下の文字起こしから話者別サマリーを作成:\n\n{transcript}",
        max_tokens=4096,
    )


def _split_transcript(text: str, max_chars: int) -> list[str]:
    lines = text.split("\n")
    chunks, current, current_len = [], [], 0
    for line in lines:
        if current_len + len(line) > max_chars and current:
            chunks.append("\n".join(current))
            current, current_len = [], 0
        current.append(line)
        current_len += len(line)
    if current and any(line for line in current):
        chunks.append("\n".join(current))
    return chunks


# ── Word文書生成 ──────────────────────────────────────────────────────
def create_docx(minutes_md: str, title: str) -> Path:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "游ゴシック"
    style.font.size = Pt(10.5)

    current_table = [None]  # mutableにするためリストで管理

    for line in minutes_md.split("\n"):
        line = line.rstrip()
        if line.startswith("# "):
            p = doc.add_heading(line[2:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            current_table[0] = None
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
            current_table[0] = None
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
            current_table[0] = None
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
            current_table[0] = None
        elif line.startswith("|"):
            cells = [c.strip() for c in line.strip("|").split("|")]
            if all(set(c) <= set("-: ") for c in cells):
                continue
            if current_table[0] is None:
                current_table[0] = doc.add_table(rows=0, cols=len(cells))
                current_table[0].style = "Table Grid"
            row = current_table[0].add_row()
            for i, text in enumerate(cells):
                if i < len(row.cells):
                    row.cells[i].text = text
        elif line.startswith("---"):
            doc.add_paragraph("─" * 40)
            current_table[0] = None
        elif line.strip():
            doc.add_paragraph(line)
            current_table[0] = None

    out_path = OUTPUT_DIR / f"{title}.docx"
    doc.save(str(out_path))
    return out_path


# ── バックグラウンドジョブ ────────────────────────────────────────────
async def _run_job(job_id: str, tmp_path: Path, do_diarize: bool, template_text: str, original_filename: str):
    def cb(pct: int, label: str):
        if job_id in _jobs:
            _jobs[job_id].update({"pct": pct, "label": label})

    try:
        loop = asyncio.get_event_loop()
        lang = "unknown"

        if do_diarize:
            cb(10, "文字起こし中（Whisper）...")
            whisper_segs, lang = await loop.run_in_executor(
                None, lambda: transcribe_with_segments(tmp_path, progress_cb=cb)
            )
            cb(60, "話者を識別中（pyannote）...")
            diarize_segs = await loop.run_in_executor(None, diarize, tmp_path)
            cb(72, "話者ラベルを整理中...")
            labeled = _assign_speaker_labels(whisper_segs, diarize_segs)
            transcript = _format_speaker_transcript(labeled)
        else:
            cb(10, "文字起こし中（Whisper）...")
            segs, lang = await loop.run_in_executor(
                None, lambda: transcribe_with_segments(tmp_path, progress_cb=cb)
            )
            lines = []
            for seg in segs:
                ts = f"[{int(seg['start']//60):02d}:{int(seg['start']%60):02d}]"
                lines.append(f"{ts} {seg['text']}")
            transcript = "\n".join(lines)

        cb(75, "議事録を生成中（Claude）...")
        minutes = await loop.run_in_executor(
            None, lambda: generate_minutes(transcript, template_text, progress_cb=cb)
        )

        cb(96, "ファイルを保存中...")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        stem = Path(original_filename).stem
        title = f"{stem}_{timestamp}"
        (OUTPUT_DIR / f"{title}.md").write_text(minutes, encoding="utf-8")
        create_docx(minutes, title)

        _jobs[job_id].update({
            "pct": 100,
            "label": "完了！",
            "done": True,
            "result": {
                "success": True,
                "language": lang,
                "transcript": transcript,
                "minutes": minutes,
                "diarization_used": do_diarize,
                "template_used": bool(template_text),
                "files": {
                    "markdown": f"/download/md/{title}",
                    "word": f"/download/docx/{title}",
                },
            },
        })
    except Exception as e:
        _jobs[job_id] = {"pct": 0, "label": "エラー", "done": True, "error": str(e)}
    finally:
        try:
            if tmp_path.exists():
                tmp_path.unlink()
        except Exception:
            pass


# ── APIエンドポイント ─────────────────────────────────────────────────
@app.post("/transcribe-start")
async def transcribe_start(
    file: UploadFile = File(...),
    use_diarization: str = Form("false"),
    template_file: UploadFile = File(None),
):
    allowed = {".mp3", ".mp4", ".m4a", ".wav", ".ogg", ".flac", ".webm", ".aac"}
    suffix = Path(file.filename).suffix.lower()
    if suffix not in allowed:
        raise HTTPException(400, f"対応形式: {', '.join(allowed)}")

    template_text = ""
    if template_file and template_file.filename:
        raw = await template_file.read()
        template_text = raw.decode("utf-8", errors="ignore")

    do_diarize = use_diarization.lower() == "true"
    job_id = str(uuid.uuid4())
    tmp_path = Path(tempfile.mktemp(suffix=suffix))
    async with aiofiles.open(tmp_path, "wb") as f:
        content = await file.read()
        await f.write(content)

    _jobs[job_id] = {"pct": 10, "label": "ファイルを受信しました...", "done": False}
    asyncio.create_task(_run_job(job_id, tmp_path, do_diarize, template_text, file.filename))

    return JSONResponse({"job_id": job_id})


@app.get("/progress/{job_id}")
async def get_progress(job_id: str):
    async def event_gen():
        try:
            while True:
                job = _jobs.get(job_id)
                if not job:
                    yield f"data: {json.dumps({'error': 'ジョブが見つかりません', 'done': True})}\n\n"
                    return
                yield f"data: {json.dumps(job)}\n\n"
                if job.get("done"):
                    _jobs.pop(job_id, None)
                    return
                await asyncio.sleep(0.3)
        except (GeneratorExit, asyncio.CancelledError):
            pass

    return StreamingResponse(
        event_gen(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no", "Connection": "keep-alive"},
    )


@app.post("/transcribe")
async def transcribe_endpoint(
    file: UploadFile = File(...),
    use_diarization: str = Form("false"),
    template_file: UploadFile = File(None),
):
    """音声ファイルをアップロードして議事録を生成する。"""
    allowed = {".mp3", ".mp4", ".m4a", ".wav", ".ogg", ".flac", ".webm", ".aac"}
    suffix = Path(file.filename).suffix.lower()
    if suffix not in allowed:
        raise HTTPException(400, f"対応形式: {', '.join(allowed)}")

    # テンプレート読み込み
    template_text = ""
    if template_file and template_file.filename:
        raw = await template_file.read()
        template_text = raw.decode("utf-8", errors="ignore")

    do_diarize = use_diarization.lower() == "true"
    tmp_path = Path(tempfile.mktemp(suffix=suffix))
    try:
        async with aiofiles.open(tmp_path, "wb") as f:
            content = await file.read()
            await f.write(content)

        loop = asyncio.get_event_loop()

        if do_diarize:
            # 話者識別モード
            whisper_segs, lang = await loop.run_in_executor(
                None, transcribe_with_segments, tmp_path
            )
            diarize_segs = await loop.run_in_executor(
                None, diarize, tmp_path
            )
            labeled = _assign_speaker_labels(whisper_segs, diarize_segs)
            transcript = _format_speaker_transcript(labeled)
        else:
            transcript, lang = await loop.run_in_executor(
                None, transcribe_only, tmp_path
            )

        minutes = await loop.run_in_executor(None, generate_minutes, transcript, template_text)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        title = f"{Path(file.filename).stem}_{timestamp}"

        (OUTPUT_DIR / f"{title}.md").write_text(minutes, encoding="utf-8")
        docx_path = create_docx(minutes, title)

        return JSONResponse({
            "success": True,
            "language": lang,
            "transcript": transcript,
            "minutes": minutes,
            "diarization_used": do_diarize,
            "template_used": bool(template_text),
            "files": {
                "markdown": f"/download/md/{title}",
                "word": f"/download/docx/{title}",
            },
        })
    finally:
        try:
            if tmp_path.exists():
                tmp_path.unlink()
        except Exception:
            pass


@app.get("/download/md/{title}")
async def download_md(title: str):
    path = OUTPUT_DIR / f"{title}.md"
    if not path.exists():
        raise HTTPException(404)
    return FileResponse(path, media_type="text/markdown", filename=f"{title}.md")


@app.get("/download/docx/{title}")
async def download_docx(title: str):
    path = OUTPUT_DIR / f"{title}.docx"
    if not path.exists():
        raise HTTPException(404)
    return FileResponse(
        path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"{title}.docx",
    )


class SpeakerProcessRequest(BaseModel):
    transcript: str
    minutes: str
    speaker_map: dict[str, str]
    title: str = ""


@app.post("/process-speakers")
async def process_speakers(req: SpeakerProcessRequest):
    """
    話者名を適用し、話者別サマリーを生成する。
    speaker_map: {"話者1": "田中", "話者2": "佐藤", ...}
    """
    loop = asyncio.get_event_loop()

    # 名前置換
    renamed_transcript = apply_speaker_names(req.transcript, req.speaker_map)
    renamed_minutes = apply_speaker_names(req.minutes, req.speaker_map)

    # 話者別サマリー（Claudeで生成）
    speaker_summary = await loop.run_in_executor(
        None, generate_speaker_summaries, renamed_transcript
    )

    # ファイル保存（タイトルがあれば上書き）
    title = req.title or datetime.now().strftime("speakers_%Y%m%d_%H%M%S")
    combined_md = f"{renamed_minutes}\n\n---\n\n{speaker_summary}"
    (OUTPUT_DIR / f"{title}.md").write_text(combined_md, encoding="utf-8")
    create_docx(combined_md, title)

    return JSONResponse({
        "success": True,
        "transcript": renamed_transcript,
        "minutes": renamed_minutes,
        "speaker_summary": speaker_summary,
        "files": {
            "markdown": f"/download/md/{title}",
            "word": f"/download/docx/{title}",
        },
    })


@app.get("/diarization-available")
async def diarization_available():
    """HF_TOKENが設定されているか確認するエンドポイント。"""
    return {"available": bool(HF_TOKEN)}


class AutoNameRequest(BaseModel):
    transcript: str


@app.post("/identify-speakers")
async def identify_speakers(req: AutoNameRequest):
    """文脈からAIが話者名を自動推定するエンドポイント。"""
    if not _has_api_key():
        raise HTTPException(status_code=400, detail="AIプロバイダーのAPIキーが設定されていません")
    speaker_map = auto_identify_speakers(req.transcript)
    return {"speaker_map": speaker_map}


# ── フロントエンド ────────────────────────────────────────────────────
@app.get("/", response_class=HTMLResponse)
async def index():
    return HTML_CONTENT


HTML_CONTENT = r"""<!DOCTYPE html>
<html lang="ja">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>議事録自動生成</title>
<style>
  * { box-sizing: border-box; margin: 0; padding: 0; }
  body { font-family: 'Noto Sans JP', 'Yu Gothic', sans-serif; background: #f5f7fa; color: #333; min-height: 100vh; }
  header { background: #1e3a5f; color: white; padding: 1.2rem 2rem; display: flex; align-items: center; gap: 1rem; }
  header h1 { font-size: 1.4rem; font-weight: 600; }
  .badge { background: #3b82f6; font-size: 0.7rem; padding: 2px 8px; border-radius: 99px; }
  main { max-width: 980px; margin: 2rem auto; padding: 0 1rem; }

  .card { background: white; border-radius: 12px; box-shadow: 0 2px 8px rgba(0,0,0,.08); padding: 2rem; margin-bottom: 1.5rem; }
  .card h2 { font-size: 1rem; font-weight: 600; color: #1e3a5f; margin-bottom: 1rem; }

  .drop-zone { border: 2px dashed #93c5fd; border-radius: 10px; padding: 3rem 2rem; text-align: center; cursor: pointer; transition: all .2s; background: #eff6ff; }
  .drop-zone:hover, .drop-zone.drag-over { border-color: #3b82f6; background: #dbeafe; }
  .drop-zone .icon { font-size: 2.5rem; margin-bottom: .5rem; }
  .drop-zone p { color: #6b7280; font-size: .9rem; }
  .drop-zone strong { color: #3b82f6; }
  #file-input { display: none; }

  .option-row { display: flex; align-items: center; gap: .8rem; margin: 1rem 0; padding: .8rem 1rem; background: #f8fafc; border-radius: 8px; border: 1px solid #e2e8f0; }
  .toggle { position: relative; width: 44px; height: 24px; flex-shrink: 0; }
  .toggle input { opacity: 0; width: 0; height: 0; }
  .slider { position: absolute; inset: 0; background: #cbd5e1; border-radius: 99px; cursor: pointer; transition: .2s; }
  .slider:before { content:''; position:absolute; width:18px; height:18px; left:3px; top:3px; background:white; border-radius:50%; transition:.2s; }
  input:checked + .slider { background: #1e3a5f; }
  input:checked + .slider:before { transform: translateX(20px); }
  .toggle-label { font-size: .9rem; flex: 1; }
  .toggle-label small { color: #64748b; font-size: .8rem; display: block; }
  .hf-warn { color: #b45309; background: #fffbeb; border: 1px solid #fde68a; border-radius: 6px; padding: .5rem .8rem; font-size: .8rem; display: none; margin-top: .5rem; }

  .btn { display: inline-flex; align-items: center; gap: .5rem; padding: .6rem 1.5rem; border-radius: 8px; border: none; cursor: pointer; font-size: .9rem; font-weight: 500; transition: all .15s; }
  .btn-primary { background: #1e3a5f; color: white; }
  .btn-primary:hover:not(:disabled) { background: #2d5a8f; }
  .btn-primary:disabled { background: #94a3b8; cursor: not-allowed; }
  .btn-secondary { background: #e2e8f0; color: #334155; }
  .btn-secondary:hover { background: #cbd5e1; }
  .btn-accent { background: #7c3aed; color: white; }
  .btn-accent:hover:not(:disabled) { background: #6d28d9; }
  .btn-accent:disabled { background: #c4b5fd; cursor: not-allowed; }

  .progress-wrap { display: none; margin-top: 1rem; }
  .progress-bar { height: 6px; background: #e2e8f0; border-radius: 99px; overflow: hidden; }
  .progress-fill { height: 100%; width: 0; background: linear-gradient(90deg, #3b82f6, #60a5fa); border-radius: 99px; transition: width .3s; animation: pulse 1.5s infinite; }
  @keyframes pulse { 0%,100%{opacity:1} 50%{opacity:.6} }
  .status-text { font-size: .85rem; color: #64748b; margin-top: .4rem; }

  /* 話者名登録パネル */
  .speaker-name-panel { display: none; background: #f0f7ff; border: 1px solid #bfdbfe; border-radius: 10px; padding: 1.2rem 1.5rem; margin-bottom: 1rem; }
  .speaker-name-panel h3 { font-size: .95rem; font-weight: 600; color: #1e3a5f; margin-bottom: .8rem; }
  .speaker-name-grid { display: grid; grid-template-columns: repeat(auto-fill, minmax(220px, 1fr)); gap: .6rem; margin-bottom: 1rem; }
  .speaker-name-row { display: flex; align-items: center; gap: .5rem; }
  .speaker-tag-label { font-size: .82rem; font-weight: 600; padding: 2px 8px; border-radius: 4px; white-space: nowrap; }
  .speaker-name-input { flex: 1; padding: .35rem .6rem; border: 1px solid #cbd5e1; border-radius: 6px; font-size: .85rem; }
  .speaker-name-input:focus { outline: none; border-color: #3b82f6; }
  .speaker-summary-note { font-size: .8rem; color: #64748b; margin-bottom: .8rem; }

  .result-section { display: none; }
  .tabs { display: flex; gap: .5rem; margin-bottom: 1rem; border-bottom: 2px solid #e2e8f0; flex-wrap: wrap; }
  .tab { padding: .5rem 1.2rem; cursor: pointer; border-radius: 6px 6px 0 0; font-size: .9rem; color: #64748b; border-bottom: 2px solid transparent; margin-bottom: -2px; }
  .tab.active { color: #1e3a5f; border-bottom-color: #1e3a5f; font-weight: 600; }
  .tab-content { display: none; }
  .tab-content.active { display: block; }

  pre { background: #f8fafc; border: 1px solid #e2e8f0; border-radius: 8px; padding: 1rem; font-size: .85rem; line-height: 1.8; white-space: pre-wrap; word-break: break-word; max-height: 500px; overflow-y: auto; }
  .speaker-line { display: flex; gap: .5rem; margin-bottom: .3rem; font-size: .88rem; line-height: 1.6; align-items: baseline; }
  .speaker-tag { font-weight: 600; white-space: nowrap; padding: 0 6px; border-radius: 4px; font-size: .8rem; }
  .s0 { background: #dbeafe; color: #1d4ed8; }
  .s1 { background: #dcfce7; color: #166534; }
  .s2 { background: #fef9c3; color: #854d0e; }
  .s3 { background: #fce7f3; color: #9d174d; }
  .s4 { background: #ede9fe; color: #5b21b6; }
  .timestamp { color: #94a3b8; font-size: .78rem; white-space: nowrap; }

  .md-display { font-size: .9rem; line-height: 1.8; max-height: 500px; overflow-y: auto; padding: 1rem; background: #fafafa; border: 1px solid #e2e8f0; border-radius: 8px; }
  .md-display h1 { font-size: 1.2rem; margin: .5rem 0 1rem; color: #1e3a5f; text-align: center; }
  .md-display h2 { font-size: 1rem; margin: 1.2rem 0 .5rem; color: #1e3a5f; border-bottom: 1px solid #e2e8f0; padding-bottom: .3rem; }
  .md-display h3 { font-size: .95rem; margin: .8rem 0 .4rem; color: #374151; }
  .md-display ul { padding-left: 1.5rem; }
  .md-display li { margin-bottom: .3rem; }
  .md-display table { border-collapse: collapse; width: 100%; margin: .5rem 0; }
  .md-display td, .md-display th { border: 1px solid #e2e8f0; padding: .4rem .7rem; font-size: .85rem; }
  .md-display th { background: #f1f5f9; font-weight: 600; }
  .md-display hr { border: none; border-top: 1px solid #e2e8f0; margin: 1rem 0; }
  /* 話者別サマリー色分け */
  .md-display .speaker-section { border-left: 3px solid #93c5fd; padding-left: .8rem; margin: .5rem 0 1rem; }

  .dl-buttons { display: flex; gap: .7rem; margin-top: 1rem; flex-wrap: wrap; }
  .info-pill { display: inline-flex; align-items: center; gap: .3rem; font-size: .8rem; background: #f0fdf4; color: #166534; border: 1px solid #bbf7d0; padding: .2rem .7rem; border-radius: 99px; margin: 0 .3rem .5rem 0; }
  .error-msg { color: #dc2626; background: #fef2f2; border: 1px solid #fecaca; border-radius: 8px; padding: .8rem 1rem; font-size: .9rem; margin-top: .8rem; }
</style>
</head>
<body>
<header>
  <h1>🎙️ 議事録自動生成</h1>
  <span class="badge">Whisper + Claude</span>
</header>
<main>
  <!-- アップロードカード -->
  <div class="card">
    <h2>音声ファイルをアップロード</h2>
    <div class="drop-zone" id="drop-zone" onclick="document.getElementById('file-input').click()">
      <div class="icon">🎵</div>
      <p><strong>クリックまたはドラッグ＆ドロップ</strong></p>
      <p>MP3 / MP4 / M4A / WAV / OGG / FLAC / WebM / AAC</p>
      <p style="margin-top:.5rem;font-size:.8rem;color:#9ca3af">長時間録音も対応（自動チャンク処理）</p>
    </div>
    <input type="file" id="file-input" accept=".mp3,.mp4,.m4a,.wav,.ogg,.flac,.webm,.aac">

    <!-- テンプレートアップロード -->
    <div class="option-row" style="flex-direction:column;align-items:flex-start;gap:.6rem">
      <div style="display:flex;align-items:center;gap:.6rem;width:100%">
        <span style="font-size:.9rem">📋 議事録テンプレート（任意）</span>
        <label class="btn btn-secondary" style="padding:.3rem .9rem;cursor:pointer;font-size:.82rem">
          ファイルを選択
          <input type="file" id="template-input" accept=".md,.txt,.text" style="display:none" onchange="onTemplateSelected()">
        </label>
        <span id="template-name" style="font-size:.82rem;color:#64748b"></span>
        <button id="template-clear" onclick="clearTemplate()" style="display:none;background:none;border:none;color:#94a3b8;cursor:pointer;font-size:.85rem">✕</button>
      </div>
      <div style="font-size:.78rem;color:#94a3b8">Markdown または テキストファイル。アップロードしたフォーマットに従って議事録を生成します。</div>
    </div>

    <div class="option-row">
      <label class="toggle">
        <input type="checkbox" id="diarize-toggle" onchange="onDiarizeToggle()">
        <span class="slider"></span>
      </label>
      <div class="toggle-label">
        👥 話者識別を有効にする
        <small>誰が話したかをラベル付け（HuggingFace Token が必要）</small>
      </div>
    </div>
    <div class="hf-warn" id="hf-warn">
      ⚠️ サーバーに HF_TOKEN が設定されていません。<br>
      <a href="https://huggingface.co/settings/tokens" target="_blank">HuggingFace</a> でトークンを取得し、<code>.env</code> に <code>HF_TOKEN=xxx</code> を設定してください。<br>
      また <a href="https://huggingface.co/pyannote/speaker-diarization-3.1" target="_blank">pyannote/speaker-diarization-3.1</a> の利用規約に同意が必要です。
    </div>

    <div style="margin-top:1rem;display:flex;gap:.7rem;align-items:center;flex-wrap:wrap">
      <button class="btn btn-primary" id="start-btn" disabled onclick="startProcess()">🚀 議事録を生成</button>
      <span id="file-name" style="font-size:.85rem;color:#64748b"></span>
    </div>
    <div class="progress-wrap" id="progress-wrap">
      <div class="progress-bar"><div class="progress-fill" id="progress-fill"></div></div>
      <p class="status-text" id="status-text">処理中...</p>
    </div>
    <div id="error-area"></div>
  </div>

  <!-- 話者名登録パネル（話者識別後に表示） -->
  <div class="card" id="speaker-name-card" style="display:none">
    <h2>👥 話者名を登録（任意）</h2>
    <p class="speaker-summary-note">識別された話者に名前をつけると、議事録・サマリーに反映されます。空欄のままでも構いません。</p>
    <div class="speaker-name-grid" id="speaker-name-grid"></div>
    <div style="display:flex;gap:.7rem;flex-wrap:wrap;align-items:center">
      <button class="btn" id="auto-name-btn" onclick="autoIdentifySpeakers()" style="background:#6c5ce7;color:#fff">
        🤖 AIが自動で名前を推定
      </button>
      <button class="btn btn-accent" id="apply-names-btn" onclick="applyNamesAndSummarize()">
        ✨ 名前を適用して話者別サマリーを生成
      </button>
      <div class="progress-wrap" id="summary-progress-wrap" style="display:inline-block;min-width:200px">
        <div class="progress-bar"><div class="progress-fill" id="summary-progress-fill" style="width:0"></div></div>
        <p class="status-text" id="summary-status-text"></p>
      </div>
    </div>
    <div id="summary-error"></div>
  </div>

  <!-- 結果カード -->
  <div class="card result-section" id="result-section">
    <h2>生成結果</h2>
    <div id="pills"></div>
    <div class="tabs">
      <div class="tab active" onclick="switchTab('minutes')">📄 議事録</div>
      <div class="tab" onclick="switchTab('transcript')">📝 文字起こし</div>
      <div class="tab" id="tab-summary-btn" onclick="switchTab('summary')" style="display:none">👤 話者別サマリー</div>
    </div>
    <div id="tab-minutes" class="tab-content active">
      <div class="md-display" id="minutes-display"></div>
    </div>
    <div id="tab-transcript" class="tab-content">
      <div id="transcript-display"></div>
    </div>
    <div id="tab-summary" class="tab-content">
      <div class="md-display" id="summary-display"></div>
    </div>
    <div class="dl-buttons">
      <a id="dl-md" class="btn btn-secondary" download>⬇ Markdownをダウンロード</a>
      <a id="dl-docx" class="btn btn-primary" download>⬇ Word文書をダウンロード</a>
    </div>
  </div>
</main>

<script>
let selectedFile = null;
let diarizationAvailable = false;
let currentEventSource = null;
let currentTitle = '';
let currentTranscript = '';
let currentMinutes = '';
let isDiarized = false;

fetch('/diarization-available')
  .then(r => r.json())
  .then(d => { diarizationAvailable = d.available; });

const dropZone = document.getElementById('drop-zone');
const fileInput = document.getElementById('file-input');
const startBtn = document.getElementById('start-btn');
const fileNameEl = document.getElementById('file-name');

dropZone.addEventListener('dragover', e => { e.preventDefault(); dropZone.classList.add('drag-over'); });
dropZone.addEventListener('dragleave', () => dropZone.classList.remove('drag-over'));
dropZone.addEventListener('drop', e => {
  e.preventDefault(); dropZone.classList.remove('drag-over');
  if (e.dataTransfer.files[0]) setFile(e.dataTransfer.files[0]);
});
fileInput.addEventListener('change', () => { if (fileInput.files[0]) setFile(fileInput.files[0]); });

function setFile(f) {
  selectedFile = f;
  fileNameEl.textContent = `${f.name} (${(f.size/1024/1024).toFixed(1)} MB)`;
  startBtn.disabled = false;
}

function onTemplateSelected() {
  const input = document.getElementById('template-input');
  const f = input.files[0];
  if (f) {
    document.getElementById('template-name').textContent = f.name;
    document.getElementById('template-clear').style.display = 'inline';
  }
}

function clearTemplate() {
  document.getElementById('template-input').value = '';
  document.getElementById('template-name').textContent = '';
  document.getElementById('template-clear').style.display = 'none';
}

function onDiarizeToggle() {
  const checked = document.getElementById('diarize-toggle').checked;
  document.getElementById('hf-warn').style.display = (checked && !diarizationAvailable) ? 'block' : 'none';
}

const TABS = ['minutes', 'transcript', 'summary'];
function switchTab(name) {
  document.querySelectorAll('.tab').forEach((t, i) =>
    t.classList.toggle('active', TABS[i] === name));
  TABS.forEach(n => {
    const el = document.getElementById('tab-' + n);
    if (el) el.classList.toggle('active', n === name);
  });
}

function setProgress(pct, text) {
  document.getElementById('progress-fill').style.width = pct + '%';
  document.getElementById('status-text').textContent = text;
}

async function startProcess() {
  if (!selectedFile) return;
  startBtn.disabled = true;
  document.getElementById('error-area').innerHTML = '';
  document.getElementById('result-section').style.display = 'none';
  document.getElementById('speaker-name-card').style.display = 'none';
  document.getElementById('progress-wrap').style.display = 'block';

  const useDiarize = document.getElementById('diarize-toggle').checked;
  isDiarized = useDiarize;
  setProgress(0, 'アップロード中...');

  try {
    const form = new FormData();
    form.append('file', selectedFile);
    form.append('use_diarization', useDiarize ? 'true' : 'false');
    const templateInput = document.getElementById('template-input');
    if (templateInput.files[0]) form.append('template_file', templateInput.files[0]);

    // Step 1: XHRでアップロード（進捗 0→10%）
    const startResp = await new Promise((resolve, reject) => {
      const xhr = new XMLHttpRequest();
      xhr.open('POST', '/transcribe-start');
      xhr.upload.onprogress = e => {
        if (e.lengthComputable) {
          const pct = Math.round((e.loaded / e.total) * 10);
          setProgress(pct, `アップロード中... ${(e.loaded/1024/1024).toFixed(1)}MB / ${(e.total/1024/1024).toFixed(1)}MB`);
        }
      };
      xhr.onload = () => {
        try { resolve(JSON.parse(xhr.responseText)); }
        catch(err) { reject(new Error(xhr.responseText || '処理に失敗しました')); }
      };
      xhr.onerror = () => reject(new Error('ネットワークエラー'));
      xhr.send(form);
    });

    if (!startResp.job_id) throw new Error(startResp.detail || 'ジョブの開始に失敗しました');
    setProgress(10, 'ファイルを受信しました...');

    // Step 2: SSEで進捗をリアルタイム受信
    const data = await new Promise((resolve, reject) => {
      if (currentEventSource) { currentEventSource.close(); currentEventSource = null; }
      currentEventSource = new EventSource('/progress/' + startResp.job_id);

      currentEventSource.onmessage = e => {
        let msg;
        try { msg = JSON.parse(e.data); } catch { return; }

        if (msg.error && !msg.result) {
          currentEventSource.close(); currentEventSource = null;
          reject(new Error(msg.error));
          return;
        }
        if (msg.pct !== undefined) setProgress(msg.pct, msg.label || '処理中...');
        if (msg.done) {
          currentEventSource.close(); currentEventSource = null;
          if (msg.result) resolve(msg.result);
          else reject(new Error(msg.error || '結果が取得できませんでした'));
        }
      };

      currentEventSource.onerror = () => {
        if (currentEventSource) { currentEventSource.close(); currentEventSource = null; }
        reject(new Error('サーバーとの接続が切れました'));
      };
    });

    if (!data.success) throw new Error(data.detail || '処理に失敗しました');

    setProgress(100, '完了！');
    setTimeout(() => { document.getElementById('progress-wrap').style.display = 'none'; }, 800);

    currentTitle = data.files.markdown.split('/').pop().replace('.md','');
    currentTranscript = data.transcript;
    currentMinutes = data.minutes;

    renderResults(data);

    if (data.diarization_used) {
      showSpeakerNamePanel(data.transcript);
    }
  } catch(e) {
    if (currentEventSource) { currentEventSource.close(); currentEventSource = null; }
    document.getElementById('progress-wrap').style.display = 'none';
    document.getElementById('error-area').innerHTML = `<p class="error-msg">❌ ${escHtml(e.message)}</p>`;
  } finally {
    startBtn.disabled = false;
  }
}

function renderResults(data) {
  let pills = `<span class="info-pill">🌐 検出言語: ${data.language || '—'}</span>`;
  if (data.diarization_used) pills += `<span class="info-pill">👥 話者識別: ON</span>`;
  if (data.template_used) pills += `<span class="info-pill">📋 テンプレート適用</span>`;
  document.getElementById('pills').innerHTML = pills;

  document.getElementById('minutes-display').innerHTML = markdownToHtml(data.minutes);

  const transcriptEl = document.getElementById('transcript-display');
  if (data.diarization_used) {
    transcriptEl.innerHTML = renderSpeakerTranscript(data.transcript);
  } else {
    transcriptEl.innerHTML = `<pre>${escHtml(data.transcript)}</pre>`;
  }

  document.getElementById('dl-md').href = data.files.markdown;
  document.getElementById('dl-docx').href = data.files.word;
  document.getElementById('result-section').style.display = 'block';
}

// ── 話者名登録パネル ─────────────────────────────────────────────────
const SPEAKER_COLORS = ['s0','s1','s2','s3','s4'];
const speakerColorMap = {};
let colorIdx = 0;

function getSpeakerColor(label) {
  if (!speakerColorMap[label]) {
    speakerColorMap[label] = SPEAKER_COLORS[colorIdx % SPEAKER_COLORS.length];
    colorIdx++;
  }
  return speakerColorMap[label];
}

function extractSpeakers(transcript) {
  const seen = new Map();
  for (const m of transcript.matchAll(/話者[^\s:：]+/g)) {
    if (!seen.has(m[0])) seen.set(m[0], null);
  }
  return [...seen.keys()];
}

function showSpeakerNamePanel(transcript) {
  // colorMapをリセット
  Object.keys(speakerColorMap).forEach(k => delete speakerColorMap[k]);
  colorIdx = 0;

  const speakers = extractSpeakers(transcript);
  const grid = document.getElementById('speaker-name-grid');
  grid.innerHTML = '';
  for (const sp of speakers) {
    const cls = getSpeakerColor(sp);
    const row = document.createElement('div');
    row.className = 'speaker-name-row';
    row.innerHTML = `
      <span class="speaker-tag-label speaker-tag ${cls}">${escHtml(sp)}</span>
      <input class="speaker-name-input" type="text" placeholder="例: 田中" data-speaker="${escHtml(sp)}" maxlength="30">
    `;
    grid.appendChild(row);
  }
  document.getElementById('speaker-name-card').style.display = 'block';
  document.getElementById('summary-progress-wrap').style.display = 'none';
}

async function autoIdentifySpeakers() {
  const transcript = currentTranscript;
  if (!transcript.trim()) {
    alert('先に文字起こしを実行してください。');
    return;
  }
  const btn = document.getElementById('auto-name-btn');
  btn.disabled = true;
  btn.textContent = '🤖 推定中...';
  try {
    const res = await fetch('/identify-speakers', {
      method: 'POST',
      headers: {'Content-Type': 'application/json'},
      body: JSON.stringify({transcript})
    });
    if (!res.ok) throw new Error(await res.text());
    const data = await res.json();
    const map = data.speaker_map || {};
    if (Object.keys(map).length === 0) {
      alert('会話の文脈から話者名を特定できませんでした。手動で入力してください。');
    } else {
      // 入力欄にセット（ユーザーが後から編集可能）
      document.querySelectorAll('.speaker-name-input').forEach(input => {
        const label = input.dataset.speaker;
        if (map[label]) input.value = map[label];
      });
      alert('AIが推定した名前をセットしました。必要に応じて修正してください。');
    }
  } catch(e) {
    alert('自動命名に失敗しました: ' + e.message);
  } finally {
    btn.disabled = false;
    btn.textContent = '🤖 AIが自動で名前を推定';
  }
}

async function applyNamesAndSummarize() {
  const inputs = document.querySelectorAll('.speaker-name-input');
  const speakerMap = {};
  for (const inp of inputs) {
    const name = inp.value.trim();
    if (name) speakerMap[inp.dataset.speaker] = name;
  }

  const btn = document.getElementById('apply-names-btn');
  btn.disabled = true;
  const pw = document.getElementById('summary-progress-wrap');
  pw.style.display = 'inline-block';
  document.getElementById('summary-error').innerHTML = '';

  let pct = 10;
  const summaryFill = document.getElementById('summary-progress-fill');
  const summaryStatus = document.getElementById('summary-status-text');
  summaryFill.style.width = '10%';
  summaryStatus.textContent = '話者名を適用中...';
  const iv = setInterval(() => {
    pct += Math.random() * 3;
    if (pct > 85) pct = 85;
    summaryFill.style.width = pct + '%';
    if (pct > 40) summaryStatus.textContent = '話者別サマリーを生成中（Claude）...';
  }, 1500);

  try {
    const res = await fetch('/process-speakers', {
      method: 'POST',
      headers: {'Content-Type': 'application/json'},
      body: JSON.stringify({
        transcript: currentTranscript,
        minutes: currentMinutes,
        speaker_map: speakerMap,
        title: currentTitle,
      }),
    });
    const data = await res.json();
    clearInterval(iv);
    if (!res.ok || !data.success) throw new Error(data.detail || 'サマリー生成に失敗しました');

    summaryFill.style.width = '100%';
    summaryStatus.textContent = '完了！';
    setTimeout(() => { pw.style.display = 'none'; }, 800);

    // 結果を更新
    currentTranscript = data.transcript;
    currentMinutes = data.minutes;

    document.getElementById('minutes-display').innerHTML = markdownToHtml(data.minutes);
    document.getElementById('transcript-display').innerHTML = renderSpeakerTranscript(data.transcript);

    // 話者別サマリータブを表示
    document.getElementById('tab-summary-btn').style.display = 'inline-block';
    document.getElementById('summary-display').innerHTML = markdownToHtml(data.speaker_summary);

    document.getElementById('dl-md').href = data.files.markdown;
    document.getElementById('dl-docx').href = data.files.word;

    // サマリータブに切替
    switchTab('summary');
  } catch(e) {
    clearInterval(iv);
    pw.style.display = 'none';
    document.getElementById('summary-error').innerHTML = `<p class="error-msg">❌ ${escHtml(e.message)}</p>`;
  } finally {
    btn.disabled = false;
  }
}

// ── レンダリング ──────────────────────────────────────────────────────
function renderSpeakerTranscript(text) {
  const lines = text.split('\n').filter(Boolean);
  return lines.map(line => {
    const m = line.match(/^(\[\d+:\d+\])\s+(\S+):\s+(.+)$/);
    if (m) {
      const [, ts, speaker, content] = m;
      const cls = getSpeakerColor(speaker);
      return `<div class="speaker-line">
        <span class="timestamp">${escHtml(ts)}</span>
        <span class="speaker-tag ${cls}">${escHtml(speaker)}</span>
        <span style="flex:1">${escHtml(content)}</span>
      </div>`;
    }
    return `<div style="font-size:.85rem;color:#64748b">${escHtml(line)}</div>`;
  }).join('');
}

function escHtml(s) {
  return String(s).replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;');
}

function markdownToHtml(md) {
  return md
    .replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;')
    .replace(/^# (.+)$/gm, '<h1>$1</h1>')
    .replace(/^## (.+)$/gm, '<h2>$1</h2>')
    .replace(/^### (.+)$/gm, '<h3>$1</h3>')
    .replace(/^---$/gm, '<hr>')
    .replace(/^\| (.+)$/gm, line => {
      const cells = line.split('|').map(c=>c.trim()).filter(Boolean);
      if (cells.every(c => /^[-:]+$/.test(c))) return '';
      return '<tr>' + cells.map(c=>`<td>${c}</td>`).join('') + '</tr>';
    })
    .replace(/(<tr>[\s\S]*?<\/tr>\n?)+/g, m => `<table>${m}</table>`)
    .replace(/^- (.+)$/gm, '<li>$1</li>')
    .replace(/(<li>[\s\S]*?<\/li>\n?)+/g, m => `<ul>${m}</ul>`)
    .replace(/\*\*(.+?)\*\*/g, '<strong>$1</strong>');
}
</script>
</body>
</html>
"""
